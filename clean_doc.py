"""
Очистка DOCX-документов от mojibake и управляющих символов
с одновременным формированием отчёта об изменениях.

 ▸ Запуск без аргументов → обрабатываются все *.docx*, лежащие
   в каталоге скрипта.
 ▸ Для каждого файла выводятся:
       — исправлено символов mojibake
       — удалено управляющих символов
 ▸ Коды возврата:
       0  — обработка успешна
       2  — DOCX-файлы не найдены
"""

from __future__ import annotations

import argparse
import sys
import unicodedata
from pathlib import Path
from typing import Iterable, Tuple

from docx import Document
from ftfy import fix_text
from tqdm import tqdm

# --------------------------------------------------------------------------- #
CONTROL_CATEGORIES = {"Cf", "Cc"}


def _strip_control_chars(text: str) -> Tuple[str, int]:
    """
    Удаляет управляющие символы (категории Cf / Cc, кроме \\t\\n\\r).
    Возвращает (очищенный_текст, количество_удалённых_символов).
    """
    new_chars = []
    removed = 0
    for ch in text:
        if unicodedata.category(ch) in CONTROL_CATEGORIES and ch not in "\t\n\r":
            removed += 1
        else:
            new_chars.append(ch)
    return "".join(new_chars), removed


def _clean_string(text: str) -> Tuple[str, int, int]:
    """
    Возвращает:
        cleaned_text,
        n_mojibake_fixes (символы, изменённые ftfy),
        n_control_removed (символы, удалённые как управляющие)
    """
    if not text:
        return text, 0, 0

    fixed = fix_text(text, normalization="NFC")
    mojibake_fixes = sum(1 for a, b in zip(text, fixed) if a != b) + abs(len(text) - len(fixed))

    cleaned, control_removed = _strip_control_chars(fixed)
    return cleaned, mojibake_fixes, control_removed


def _process_paragraphs(paragraphs) -> Tuple[int, int]:
    """Обрабатывает runs абзацев. Возвращает суммарные счётчики."""
    mojibake_total = 0
    control_total = 0

    for para in paragraphs:
        for run in para.runs:
            new_text, mb, ctrl = _clean_string(run.text)
            if new_text != run.text:
                run.text = new_text
            mojibake_total += mb
            control_total += ctrl

    return mojibake_total, control_total


def clean_docx(src: Path, dst: Path) -> Tuple[int, int]:
    """
    Очищает документ и сохраняет результат.
    Возвращает (исправлено_mojibake, удалено_управляющих).
    """
    doc = Document(src)

    mb_total, ctrl_total = _process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                mb, ctrl = _process_paragraphs(cell.paragraphs)
                mb_total += mb
                ctrl_total += ctrl

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    return mb_total, ctrl_total


# -------------------------- CLI-вспомогательные функции -------------------- #
def _iter_docx(targets: Iterable[str | Path]) -> Iterable[Path]:
    """Разворачивает список аргументов в реальные *.docx*-пути."""
    for t in targets:
        p = Path(t)
        if p.is_file() and p.suffix.lower() == ".docx":
            yield p
        else:                              # glob-шаблоны типа *.docx
            yield from p.parent.glob(p.name)


def _build_argparser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(
        prog="clean_doc.py",
        description=(
            "Очистка DOCX-файлов от mojibake и управляющих символов "
            "с выводом отчёта по каждому файлу.\n"
            "Если аргументы опущены — обрабатываются все *.docx* "
            "в каталоге скрипта."
        ),
        formatter_class=argparse.RawTextHelpFormatter,
    )
    ap.add_argument(
        "targets",
        nargs="*",
        help="Файлы *.docx* или шаблоны (*.docx). "
        "Если опущено — берутся все *.docx* рядом со скриптом.",
    )
    ap.add_argument(
        "-b",
        "--backup",
        action="store_true",
        help="Создавать резервную копию *.bak* перед перезаписью.",
    )
    ap.add_argument(
        "--suffix",
        default="_clean",
        help="Суффикс очищенных файлов (по умолчанию '_clean').",
    )
    ap.add_argument(
        "-q",
        "--quiet",
        action="store_true",
        help="Подавить прогресс-бар и подробные отчёты.",
    )
    return ap


# ----------------------------------- main ---------------------------------- #
def main(argv: list[str] | None = None) -> None:
    argv = argv or sys.argv[1:]
    ap = _build_argparser()
    args = ap.parse_args(argv)

    # Если цели не заданы → работаем в папке скрипта
    if args.targets:
        candidates = list(_iter_docx(args.targets))
    else:
        script_dir = Path(__file__).resolve().parent
        candidates = list(script_dir.glob("*.docx"))

    if not candidates:
        ap.error("DOCX-файлы не найдены.")

    iterator = candidates if args.quiet else tqdm(candidates, desc="Cleaning DOCX")
    total_mb, total_ctrl = 0, 0

    for src in iterator:
        dst = src.with_stem(src.stem + args.suffix)

        if args.backup and dst.exists():
            src.replace(src.with_suffix(".bak"))

        mb, ctrl = clean_docx(src, dst)
        total_mb += mb
        total_ctrl += ctrl

        if not args.quiet:
            print(
                f"{src.name}: исправлено mojibake = {mb}, "
                f"удалено управляющих = {ctrl}"
            )

    if not args.quiet:
        print(
            f"\nОбработано файлов: {len(candidates)} | "
            f"всего исправлено символов mojibake: {total_mb} | "
            f"всего удалено управляющих: {total_ctrl}"
        )


if __name__ == "__main__":
    main()
